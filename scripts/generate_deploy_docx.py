"""Genere le guide de deploiement Render en .docx mis en forme."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# -- Couleurs charte Urban Land --
BRAND = RGBColor(0x1D, 0x4E, 0xD8)   # bleu principal
ACCENT = RGBColor(0x7C, 0x3A, 0xED)  # violet accent
AMBER = RGBColor(0xD9, 0x77, 0x06)
GREEN = RGBColor(0x05, 0x96, 0x69)
ROSE = RGBColor(0xBE, 0x12, 0x3C)
SLATE = RGBColor(0x47, 0x55, 0x69)
LIGHT_BG = "F1F5F9"
INFO_BG = "EFF6FF"
WARN_BG = "FEF3C7"
DANGER_BG = "FEE2E2"
SUCCESS_BG = "DCFCE7"


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_h(doc, text, level=1, color=BRAND, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt({1: 22, 2: 16, 3: 13, 4: 11}.get(level, 11))
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    return p


def add_p(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Light grey background
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BG)
    p_pr = p._p.get_or_add_pPr()
    p_pr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_callout(doc, label, body, color_bg, color_text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color_text
    run2 = p.add_run(body)
    run2.font.size = Pt(10)
    run2.font.color.rgb = SLATE
    doc.add_paragraph()


def add_kv_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)
    # Header
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        set_cell_bg(c, "1D4ED8")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    # Rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = table.cell(ri, ci)
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    doc.add_paragraph()


def add_step(doc, num, title, color=BRAND):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r1 = p.add_run(f"Étape {num} — ")
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(13)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.color.rgb = SLATE
    r2.font.size = Pt(13)


# =============================================================================
# Generation du document
# =============================================================================
doc = Document()
# Page setup
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

# -- Cover -------------------------------------------------------------------
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("URBAN LAND")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Solutions Patrimoniales")
r.font.size = Pt(12)
r.font.color.rgb = SLATE
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Guide complet de déploiement\nsur Render.com")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ACCENT

doc.add_paragraph()
doc.add_paragraph()
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = intro.add_run(
    "Procédure étape par étape pour mettre Urban Land en production "
    "sur Render.com avec base PostgreSQL, Redis, API Django, "
    "workers Celery et frontend React."
)
r.font.size = Pt(11)
r.font.color.rgb = SLATE
r.italic = True

doc.add_paragraph()
doc.add_paragraph()
v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run("Version 1.0  |  Production-ready")
r.font.size = Pt(9)
r.font.color.rgb = SLATE
doc.add_page_break()

# -- TOC ---------------------------------------------------------------------
add_h(doc, "Table des matières", level=1)
toc_items = [
    "1. Vue d'ensemble",
    "2. Prérequis",
    "3. Étape 1 — Préparer le dépôt GitHub",
    "4. Étape 2 — Connecter GitHub à Render",
    "5. Étape 3 — Créer le Blueprint",
    "6. Étape 4 — Suivre le premier déploiement",
    "7. Étape 5 — Récupérer les URLs",
    "8. Étape 6 — Configurer les variables d'environnement",
    "9. Étape 7 — Première connexion",
    "10. Étape 8 — Domaine personnalisé (optionnel)",
    "11. Étape 9 — Mises à jour de l'application",
    "12. Étape 10 — Sauvegardes & maintenance",
    "13. Étape 11 — Dépannage",
    "14. Étape 12 — Coûts estimés",
    "15. Checklist finale",
    "16. Ressources utiles",
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(11)
doc.add_page_break()

# -- 1. Vue d'ensemble -------------------------------------------------------
add_h(doc, "1. Vue d'ensemble", level=1)
add_p(doc,
      "Render héberge 6 services interconnectés définis dans le fichier "
      "render.yaml de votre projet :")
add_kv_table(
    doc,
    ["#", "Service", "Type", "Rôle"],
    [
        ["1", "urban-land-db", "PostgreSQL", "Base de données principale"],
        ["2", "urban-land-redis", "Redis", "Cache + broker Celery"],
        ["3", "urban-land-api", "Web (Python)", "API Django + Gunicorn"],
        ["4", "urban-land-worker", "Background Worker", "Tâches Celery"],
        ["5", "urban-land-beat", "Background Worker", "Planificateur (rappels d'échéances)"],
        ["6", "urban-land-web", "Static Site", "Frontend React (Vite build)"],
    ],
    col_widths=[1, 5, 4, 6],
)
add_callout(
    doc,
    "💡 INFO",
    "Toute l'infrastructure est définie comme code dans render.yaml. "
    "Le Blueprint Render lit ce fichier et crée automatiquement les 6 services connectés.",
    INFO_BG, BRAND,
)

# -- 2. Prérequis ------------------------------------------------------------
add_h(doc, "2. Prérequis", level=1)
add_bullet(doc, "Un compte GitHub (gratuit)")
add_bullet(doc, "Un compte Render gratuit — https://dashboard.render.com")
add_bullet(doc, "Le code Urban Land poussé sur un dépôt GitHub (public ou privé)")
add_bullet(doc, "Git installé localement (https://git-scm.com)")
add_bullet(doc, "Aucun service Render conflictuel à supprimer au préalable")

# -- 3. Étape 1 : GitHub -----------------------------------------------------
add_step(doc, 1, "Préparer le dépôt GitHub")

add_h(doc, "A. Créer le dépôt GitHub", level=3, color=ACCENT)
add_numbered(doc, "Allez sur https://github.com/new")
add_numbered(doc, "Nom du dépôt : urban-land-management-app")
add_numbered(doc, "Visibilité : Private (recommandé) ou Public")
add_numbered(doc, "NE PAS initialiser avec README ou .gitignore")
add_numbered(doc, "Cliquez Create repository")

add_h(doc, "B. Pousser votre code local", level=3, color=ACCENT)
add_p(doc, "Ouvrez PowerShell dans le dossier du projet :")
add_code(doc, """cd C:\\Users\\amani\\DataspellProjects\\urban-land-management-app

# Initialiser git si pas déjà fait
git init
git add .
git commit -m "Initial commit Urban Land"

# Connecter au dépôt GitHub
git branch -M main
git remote add origin https://github.com/<TON-USER>/urban-land-management-app.git

# Pousser
git push -u origin main""")
add_callout(
    doc,
    "⚠️ AUTHENTIFICATION GIT",
    "Si Git vous demande un mot de passe, utilisez un Personal Access Token GitHub : "
    "Settings → Developer settings → Tokens (classic) → Generate new token avec scope 'repo'.",
    WARN_BG, AMBER,
)

add_h(doc, "C. Vérifier les fichiers critiques", level=3, color=ACCENT)
add_p(doc, "Sur GitHub, confirmez la présence à la racine du dépôt :")
for fichier in [
    "render.yaml", "DEPLOY.md", "backend/build.sh",
    "backend/requirements.txt", "backend/manage.py",
    "frontend/package.json", "frontend/vite.config.ts",
]:
    add_bullet(doc, fichier)

# -- 4. Étape 2 : Connect GitHub --------------------------------------------
add_step(doc, 2, "Connecter GitHub à Render")
add_numbered(doc, "Allez sur https://dashboard.render.com")
add_numbered(doc, "Si pas encore inscrit : Sign Up avec votre compte GitHub")
add_numbered(doc, "Menu en haut à droite → Account Settings → GitHub")
add_numbered(doc, "Cliquez Configure account → autorisez Render à accéder à votre dépôt")
add_numbered(doc, "Choisissez : All repositories (recommandé) OU Only select repositories → cochez urban-land-management-app")
add_numbered(doc, "Cliquez Save")

# -- 5. Étape 3 : Blueprint -------------------------------------------------
add_step(doc, 3, "Créer le Blueprint (déploiement automatique)")
add_p(doc, "Le Blueprint lit render.yaml et crée les 6 services automatiquement.")
add_numbered(doc, "Dashboard Render → menu de gauche → Blueprints")
add_numbered(doc, "Bouton bleu New Blueprint Instance")
add_numbered(doc, "Connect a repository → sélectionnez amanid/urban-land-management-app")
add_numbered(doc, "Branch : main")
add_numbered(doc, "Blueprint Name : urban-land (ou un nom de votre choix)")
add_numbered(doc, "Render analyse render.yaml et affiche la prévisualisation des 6 services")
add_p(doc, "Vérifiez que tous les services apparaissent :", bold=True)
for s in [
    "urban-land-db (Database PostgreSQL)",
    "urban-land-redis (Redis)",
    "urban-land-api (Web Service)",
    "urban-land-worker (Background Worker)",
    "urban-land-beat (Background Worker)",
    "urban-land-web (Static Site)",
]:
    add_bullet(doc, s)
add_numbered(doc, "Cliquez Apply")
add_callout(
    doc, "⏱️ DÉPLOIEMENT EN COURS",
    "Render commence à déployer. Surveillez la progression dans l'onglet Events.",
    INFO_BG, BRAND,
)

# -- 6. Étape 4 : Suivi -----------------------------------------------------
add_step(doc, 4, "Suivre le premier déploiement")
add_p(doc, "Le déploiement prend environ 8-12 minutes la première fois. Voici l'ordre normal :")

add_h(doc, "4.1 Base de données (10 sec)", level=3)
add_bullet(doc, "urban-land-db → status Available ✅")

add_h(doc, "4.2 Redis (10 sec)", level=3)
add_bullet(doc, "urban-land-redis → status Available ✅")

add_h(doc, "4.3 API Django (~5 min)", level=3)
add_p(doc, "Surveillez les logs : urban-land-api → onglet Logs. Vous verrez :")
add_code(doc, """==> Installing Python dependencies
Successfully installed django-5.0.6 djangorestframework-3.15.1 ...
==> Collecting static files
==> Applying database migrations
Applying contenttypes.0001_initial... OK
Applying accounts.0001_initial... OK
...
==> Seeding roles & permissions
Cree: Super administrateur (super_admin)
...
==> Seeding demo users
Cree: superadmin@urban-land.local  (role: super_admin, mot de passe: <RANDOM>)
...
[INFO] Listening at: http://0.0.0.0:10000""")
add_callout(
    doc, "⚠️ IMPORTANT",
    "Notez bien le mot de passe initial affiché dans les logs (généré aléatoirement par Render). "
    "Vous en aurez besoin pour la première connexion.",
    WARN_BG, AMBER,
)

add_h(doc, "4.4 Workers Celery (~3 min chacun)", level=3)
add_bullet(doc, "urban-land-worker → Live")
add_bullet(doc, "urban-land-beat → Live")

add_h(doc, "4.5 Frontend (~3 min)", level=3)
add_p(doc, "urban-land-web :")
add_code(doc, """> urban-land-frontend@1.0.0 build
> tsc -b && vite build
✓ built in 12.34s
==> Uploading build...
==> Your site is live 🎉""")

# -- 7. Étape 5 : URLs ------------------------------------------------------
add_step(doc, 5, "Récupérer les URLs publiques")
add_p(doc, "Une fois tout déployé, récupérez les URLs publiques visibles en haut de chaque service :")
add_kv_table(
    doc,
    ["Service", "URL publique"],
    [
        ["API", "https://urban-land-api.onrender.com"],
        ["Frontend", "https://urban-land-web.onrender.com"],
        ["Admin Django", "https://urban-land-api.onrender.com/admin/"],
        ["Healthcheck", "https://urban-land-api.onrender.com/api/v1/dashboard/health/"],
    ],
    col_widths=[4, 12],
)
add_callout(
    doc, "💡 TEST RAPIDE",
    'Ouvrez l\'URL Healthcheck dans votre navigateur. Vous devez voir : {"status":"ok","service":"urban-land","version":"1.0.0"}',
    INFO_BG, BRAND,
)

# -- 8. Étape 6 : Variables ------------------------------------------------
add_step(doc, 6, "Configurer les variables d'environnement")

add_h(doc, "6.1 Variables auto-générées par Render", level=3)
add_p(doc, "Déjà configurées automatiquement :")
add_bullet(doc, "SECRET_KEY (générée aléatoirement)")
add_bullet(doc, "DATABASE_URL (lien vers urban-land-db)")
add_bullet(doc, "REDIS_URL (lien vers urban-land-redis)")
add_bullet(doc, "DEMO_USERS_PASSWORD (générée aléatoirement)")

add_h(doc, "6.2 Variables à renseigner manuellement", level=3)
add_p(doc, "Dans le service urban-land-api → onglet Environment → ajoutez :")
add_kv_table(
    doc,
    ["Variable", "Valeur d'exemple"],
    [
        ["CORS_ALLOWED_ORIGINS", "https://urban-land-web.onrender.com  ⚠️ OBLIGATOIRE"],
        ["COMPANY_NAME", "Ma Société Immobilière SARL"],
        ["COMPANY_ADDRESS", "Cocody II-Plateaux, Abidjan"],
        ["COMPANY_PHONE", "+225 27 22 49 12 34"],
        ["COMPANY_EMAIL", "contact@masociete.ci"],
        ["COMPANY_RCCM", "CI-ABJ-2026-B-12345"],
        ["COMPANY_TAX_ID", "00000000"],
        ["DEFAULT_FROM_EMAIL", "no-reply@masociete.ci"],
    ],
    col_widths=[6, 10],
)
add_p(doc, "Pour chaque variable : Environment → Add Environment Variable → Key + Value → Save Changes. "
           "Render redéploie automatiquement (~2 min).")

add_h(doc, "6.3 Variables optionnelles (email/SMS)", level=3)
add_p(doc, "Si vous voulez activer les rappels d'échéances par email :")
add_kv_table(
    doc,
    ["Variable", "Exemple"],
    [
        ["EMAIL_BACKEND", "django.core.mail.backends.smtp.EmailBackend"],
        ["EMAIL_HOST", "smtp.sendgrid.net"],
        ["EMAIL_PORT", "587"],
        ["EMAIL_HOST_USER", "apikey"],
        ["EMAIL_HOST_PASSWORD", "<votre API key SendGrid>"],
        ["EMAIL_USE_TLS", "True"],
    ],
    col_widths=[6, 10],
)

# -- 9. Étape 7 : Connexion -------------------------------------------------
add_step(doc, 7, "Première connexion")

add_h(doc, "7.1 Récupérer le mot de passe initial", level=3)
add_numbered(doc, "Service urban-land-api → onglet Environment")
add_numbered(doc, "Cherchez DEMO_USERS_PASSWORD → cliquez l'œil 👁️ pour révéler")
add_numbered(doc, "Copiez cette valeur")

add_h(doc, "7.2 Se connecter", level=3)
add_numbered(doc, "Ouvrez votre frontend : https://urban-land-web.onrender.com")
add_numbered(doc, "Connectez-vous avec : superadmin@urban-land.local + mot de passe noté")
add_numbered(doc, "Vous arrivez sur le tableau de bord Super Admin")

add_h(doc, "7.3 Changer votre mot de passe IMMÉDIATEMENT", level=3)
add_numbered(doc, "Cliquez sur votre avatar en haut à droite")
add_numbered(doc, "Mot de passe → tapez l'ancien + un nouveau (≥ 8 caractères)")
add_numbered(doc, "Enregistrer")

add_h(doc, "7.4 Configurer votre équipe", level=3)
add_numbered(doc, "Sidebar → Utilisateurs → Nouvel utilisateur")
add_numbered(doc, "Créez un compte pour chaque collaborateur avec le bon rôle")
add_numbered(doc, "Communiquez à chacun ses identifiants + mot de passe initial")

add_h(doc, "7.5 Désactiver les comptes de démo (recommandé)", level=3)
add_numbered(doc, "Service urban-land-api → Environment")
add_numbered(doc, "Modifiez DEMO_USERS_ENABLED : passez à False")
add_numbered(doc, "Sauvegardez")

# -- 10. Étape 8 : Domaine ---------------------------------------------------
add_step(doc, 8, "Domaine personnalisé (optionnel)")

add_h(doc, "Pour le frontend", level=3)
add_numbered(doc, "Service urban-land-web → Settings → Custom Domains")
add_numbered(doc, "Cliquez Add Custom Domain")
add_numbered(doc, "Tapez votre domaine : app.masociete.ci")
add_numbered(doc, "Render fournit un enregistrement DNS à créer :")
add_code(doc, """Type:   CNAME
Name:   app
Value:  urban-land-web.onrender.com""")
add_numbered(doc, "Ajoutez cet enregistrement chez votre registrar (OVH, Gandi, Cloudflare…)")
add_numbered(doc, "Attendez 5-30 min pour la propagation DNS")
add_numbered(doc, "Render génère automatiquement le certificat SSL Let's Encrypt 🔒")

add_h(doc, "Pour l'API (optionnel)", level=3)
add_p(doc, "Même procédure avec api.masociete.ci → CNAME vers urban-land-api.onrender.com.")

add_h(doc, "Après ajout du domaine personnalisé", level=3)
add_p(doc, "Mettez à jour CORS_ALLOWED_ORIGINS dans urban-land-api :")
add_code(doc, "https://app.masociete.ci,https://urban-land-web.onrender.com")

# -- 11. Étape 9 : Mises à jour ---------------------------------------------
add_step(doc, 9, "Mises à jour de l'application")
add_p(doc, "Toute modification poussée sur la branche main redéclenche automatiquement un nouveau déploiement.")
add_code(doc, """cd C:\\Users\\amani\\DataspellProjects\\urban-land-management-app

# Modifier votre code...
git add .
git commit -m "feat: nouvelle fonctionnalité"
git push origin main""")
add_p(doc, "Render rebuilde et redéploie en 2-5 minutes. Vous recevez un email de notification de succès/échec.")

add_h(doc, "Rollback en cas de problème", level=3)
add_numbered(doc, "Service → onglet Events")
add_numbered(doc, "Trouvez un déploiement antérieur qui fonctionnait")
add_numbered(doc, "Cliquez Rollback to this deploy")
add_numbered(doc, "Render restaure l'ancienne version en 30 secondes")

# -- 12. Étape 10 : Sauvegardes ---------------------------------------------
add_step(doc, 10, "Sauvegardes & maintenance")

add_h(doc, "Sauvegardes PostgreSQL", level=3)
add_bullet(doc, "Plan Free : pas de sauvegarde automatique")
add_bullet(doc, "Plan Starter+ : snapshots quotidiens automatiques (rétention 7 jours)")
add_bullet(doc, "Export manuel depuis votre ordinateur :")
add_code(doc, """# Récupérez l'URL de connexion externe depuis Render → DB → Connect → External
PGPASSWORD=<password> pg_dump -h <render-host> -U <user> -d urban_land > backup.sql""")

add_h(doc, "Monitoring", level=3)
add_bullet(doc, "Logs : onglet Logs de chaque service (temps réel)")
add_bullet(doc, "Metrics : CPU / RAM / Bande passante (onglet Metrics)")
add_bullet(doc, "Alertes : configurez des notifications Slack ou email dans Settings → Notifications")

add_h(doc, "Surveillance applicative", level=3)
add_p(doc, "Depuis Urban Land en tant que Super Admin :")
add_bullet(doc, "Tableau de bord → bandeau 'Signaux fraude' + fil d'activité temps réel")
add_bullet(doc, "Audit & intégrité → vérification hash chain SHA-256")
add_bullet(doc, "Maintenance → purge des données de test si besoin")

# -- 13. Étape 11 : Dépannage -----------------------------------------------
add_step(doc, 11, "Dépannage des erreurs fréquentes", color=ROSE)

errors = [
    ("'Service crashed during build'",
     "Dépendance manquante ou erreur Python.",
     "Vérifiez requirements.txt, consultez les logs Events, lancez 'pip install -r backend/requirements.txt' localement."),
    ("'502 Bad Gateway'",
     "L'API est en cours de réveil (plan free met en pause après 15 min d'inactivité).",
     "Patientez 30-60 secondes et rechargez."),
    ("'CORS error' dans la console du navigateur",
     "CORS_ALLOWED_ORIGINS ne contient pas l'URL du frontend.",
     "Service urban-land-api → Environment → vérifiez CORS_ALLOWED_ORIGINS = URL exacte du frontend avec https://."),
    ("'Failed to read dockerfile'",
     "Vous avez créé un service Web Service en mode Docker manuellement au lieu d'utiliser le Blueprint.",
     "Supprimez le service, recommencez avec Blueprints → New Blueprint Instance."),
    ("'Database connection refused'",
     "Transitoire (maintenance Render) ou plan free indisponible.",
     "Vérifiez le statut de la DB dans le dashboard. Augmentez le --timeout Gunicorn si besoin."),
    ("'Frontend ne se charge pas / page blanche'",
     "Erreur TypeScript ou build cassé.",
     "Service urban-land-web → Logs. Vérifiez que 'npm run build' n'a pas d'erreur."),
    ("'Demo users not seeding'",
     "DEMO_USERS_ENABLED=False ou collision sur les emails.",
     "Service urban-land-api → Environment → DEMO_USERS_ENABLED=True. Ou via Shell : 'python manage.py seed_demo_users --force'."),
]
for err, cause, solution in errors:
    add_h(doc, "🔴 " + err, level=3, color=ROSE)
    add_p(doc, "Cause : ", bold=True)
    add_p(doc, cause)
    add_p(doc, "Solution : ", bold=True)
    add_p(doc, solution)

# -- 14. Étape 12 : Coûts ---------------------------------------------------
add_step(doc, 12, "Coûts estimés", color=GREEN)
add_kv_table(
    doc,
    ["Plan", "Coût/mois", "Usage"],
    [
        ["Free (tout en free)", "0 $", "Démo uniquement. Sleep après 15 min. RAM 512MB."],
        ["Starter", "~42 $ (7$ × 5 services + DB)", "Pilote / petite production. Always-on, 512MB RAM."],
        ["Standard", "~150 $", "Production sérieuse. 2GB RAM, plus de bande passante."],
    ],
    col_widths=[4, 5, 7],
)
add_callout(
    doc, "💡 CONSEIL DÉMARRAGE",
    "Commencez en Free pour tester. Passez en Starter quand vous mettez en production "
    "(les services free se mettent en veille — gênant pour les utilisateurs réels).",
    SUCCESS_BG, GREEN,
)

add_h(doc, "Passer un service en Starter", level=3)
add_numbered(doc, "Service → Settings")
add_numbered(doc, "Instance Type → choisir Starter")
add_numbered(doc, "Save Changes")

# -- 15. Checklist ----------------------------------------------------------
add_h(doc, "Checklist finale", level=1, color=GREEN)
add_p(doc, "Après ces étapes, vérifiez que tout fonctionne :")
checks = [
    "Frontend accessible : https://urban-land-web.onrender.com charge la page de login",
    "Login OK avec superadmin et le mot de passe initial",
    "Mot de passe Super Admin changé après première connexion",
    "Variables COMPANY_* renseignées (apparaissent sur les reçus PDF)",
    "CORS configuré (CORS_ALLOWED_ORIGINS)",
    "Création de lot/client/vente fonctionne",
    "Téléchargement de PDF fonctionne",
    "Comptes des collaborateurs créés (page Utilisateurs)",
    "DEMO_USERS_ENABLED=False (production)",
    "(Optionnel) Domaine personnalisé configuré",
    "(Recommandé) Plan Starter activé pour éviter le sleep",
]
for c in checks:
    p = doc.add_paragraph()
    r = p.add_run("☐  ")
    r.font.size = Pt(11)
    r2 = p.add_run(c)
    r2.font.size = Pt(10)

# -- 16. Ressources ---------------------------------------------------------
add_h(doc, "Ressources utiles", level=1, color=ACCENT)
add_kv_table(
    doc,
    ["Ressource", "URL"],
    [
        ["Doc Render", "https://render.com/docs"],
        ["Status Render", "https://status.render.com"],
        ["Support Render", "support@render.com"],
        ["Communauté Render", "https://community.render.com"],
        ["Doc Urban Land", "DEPLOY.md à la racine du projet"],
        ["Guide intégré dans l'app", "/guide après connexion"],
    ],
    col_widths=[5, 11],
)

# -- 17. Final --------------------------------------------------------------
doc.add_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run("🚀  Bon lancement avec Urban Land !  🚀")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = BRAND

doc.add_paragraph()
sign = doc.add_paragraph()
sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sign.add_run("© 2026 Urban Land · Tous droits réservés")
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = SLATE

# -- Save -------------------------------------------------------------------
output = Path(__file__).resolve().parent.parent / "docs" / "Guide-Deploiement-Render.docx"
output.parent.mkdir(parents=True, exist_ok=True)
doc.save(output)
print(f"\n✅ Document genere : {output}")
print(f"   Taille : {output.stat().st_size / 1024:.1f} Ko")
